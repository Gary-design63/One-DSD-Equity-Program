"""
One DSD Equity Program — Document & Archive Generator
Produces:
  1. One_DSD_Equity_Program_PRD.docx  — formatted Word document
  2. One_DSD_Equity_Program.zip        — full source + docs archive
"""

import os
import re
import zipfile
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).parent.parent
OUT  = BASE / "exports"
OUT.mkdir(exist_ok=True)

# ── Color Palette ──────────────────────────────────────────────────────────────
PRIMARY  = RGBColor(0x00, 0x38, 0x65)   # #003865 — DHS dark blue
ACCENT   = RGBColor(0x78, 0xBE, 0x21)   # #78BE21 — DHS green
MUTED    = RGBColor(0x5A, 0x65, 0x77)   # #5A6577
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
BLACK    = RGBColor(0x1A, 0x1A, 0x2E)
LIGHT_BG = RGBColor(0xF8, 0xF9, 0xFA)

def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_para_border_left(para, color_hex='003865', size=24):
    """Add a left border to a paragraph (for callout boxes)."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    str(size))
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)

def add_page_break(doc):
    doc.add_page_break()

# ── Word Document Builder ──────────────────────────────────────────────────────

def build_word():
    doc = Document()

    # ── Page layout ──
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(1.1)
    section.top_margin  = section.bottom_margin = Inches(1.0)

    # ── Core styles ──
    styles = doc.styles

    def style_normal():
        s = styles['Normal']
        s.font.name     = 'Calibri'
        s.font.size     = Pt(10.5)
        s.font.color.rgb = BLACK
        s.paragraph_format.space_after = Pt(6)

    style_normal()

    def h(level, text, color=PRIMARY):
        tag = f'Heading {level}'
        p   = doc.add_heading(text, level=level)
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.color.rgb = color
        run.font.bold      = True
        if level == 1:
            run.font.size = Pt(22)
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after  = Pt(8)
        elif level == 2:
            run.font.size = Pt(16)
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after  = Pt(6)
        elif level == 3:
            run.font.size = Pt(13)
            run.font.color.rgb = PRIMARY
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)
        elif level == 4:
            run.font.size = Pt(11)
            run.font.color.rgb = MUTED
            p.paragraph_format.space_before = Pt(8)
        return p

    def body(text, bold_parts=None):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        if bold_parts:
            for part in bold_parts:
                idx = text.find(part)
                if idx >= 0:
                    run = p.runs[0]
                    run.bold = False
        return p

    def bullet(text, level=0, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
        p.paragraph_format.space_after = Pt(3)
        if bold_prefix and text.startswith(bold_prefix):
            run = p.add_run(bold_prefix)
            run.bold = True
            p.add_run(text[len(bold_prefix):])
        else:
            p.add_run(text)
        return p

    def callout(text, color_hex='003865', bg=False):
        p = doc.add_paragraph()
        set_para_border_left(p, color_hex)
        p.paragraph_format.left_indent  = Inches(0.2)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(text)
        run.font.italic = True
        run.font.color.rgb = MUTED
        return p

    def table(headers, rows, col_widths=None):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = 'Table Grid'

        # Header row
        hdr = t.rows[0]
        for i, h_text in enumerate(headers):
            cell = hdr.cells[i]
            set_cell_bg(cell, '003865')
            run = cell.paragraphs[0].add_run(h_text)
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(9.5)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Data rows
        for r_idx, row in enumerate(rows):
            tr = t.rows[r_idx + 1]
            bg = 'F8F9FA' if r_idx % 2 == 0 else 'FFFFFF'
            for c_idx, cell_text in enumerate(row):
                cell = tr.cells[c_idx]
                set_cell_bg(cell, bg)
                run = cell.paragraphs[0].add_run(str(cell_text))
                run.font.size = Pt(9)

        # Column widths
        if col_widths:
            for i, w in enumerate(col_widths):
                for row in t.rows:
                    row.cells[i].width = Inches(w)

        doc.add_paragraph()  # spacing after table
        return t

    def divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run('─' * 80)
        run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
        run.font.size = Pt(6)

    # ═══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════════════════

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('ONE DSD EQUITY PROGRAM')
    run.font.name  = 'Calibri'
    run.font.size  = Pt(28)
    run.font.bold  = True
    run.font.color.rgb = PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run('Product Requirements Document & System Instructions')
    run.font.name  = 'Calibri'
    run.font.size  = Pt(16)
    run.font.color.rgb = MUTED

    doc.add_paragraph()
    divider()
    doc.add_paragraph()

    meta_rows = [
        ('Version',         '1.0'),
        ('Date',            datetime.now().strftime('%B %Y')),
        ('Owner',           'Equity and Inclusion Operations Consultant'),
        ('Organization',    'Disability Services Division — Minnesota DHS'),
        ('Classification',  'Operational / Open Source / Non-PII'),
        ('Status',          'Active'),
    ]

    for label, value in meta_rows:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f'{label}:  ')
        r1.bold = True
        r1.font.color.rgb = PRIMARY
        r1.font.size = Pt(11)
        r2 = p.add_run(value)
        r2.font.size = Pt(11)
        r2.font.color.rgb = BLACK

    doc.add_paragraph()
    doc.add_paragraph()

    vision_p = doc.add_paragraph()
    vision_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = vision_p.add_run(
        '"A DEIA operations infrastructure so intelligent and capable that a single Consultant\n'
        'can deliver the equity support capacity of a full team — at scale, with consistency,\n'
        'and with measurable outcomes — for an entire division of 150 staff members."'
    )
    run.font.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — EXECUTIVE SUMMARY
    # ═══════════════════════════════════════════════════════════════════════════

    h(1, '1. Executive Summary')

    body(
        'The One DSD Equity Program is an AI-powered, agentic DEIA (Diversity, Equity, Inclusion, '
        'and Accessibility) operations system purpose-built for the Equity and Inclusion Operations '
        'Consultant serving the Minnesota Department of Human Services (DHS), Disability Services '
        'Division (DSD).'
    )

    callout(
        'This system is not a compliance tracker. It is an operational force multiplier — a full-scale '
        'DEIA infrastructure platform that enables one Consultant to operate at the capacity of a '
        'coordinated team, serving approximately 150 DSD staff members across all roles and levels.'
    )

    h(3, 'What This System Does')
    items = [
        'Maintains a living, structured knowledge base of all equity-relevant laws, policies, and frameworks',
        'Guides execution of 7 equity analysis workflows with stage-by-stage AI intelligence',
        'Monitors program KPIs and generates performance narratives automatically',
        'Curates and recommends DEIA learning paths for 150 staff members',
        'Tracks risks and action items with proactive daily escalation scanning',
        'Provides community engagement strategy and stakeholder intelligence',
        'Answers any equity policy, process, or practice question in real time',
        'Operates 7 AI agents semi-autonomously and in coordination with each other',
    ]
    for item in items:
        bullet(item)

    h(3, 'Design Constraints')
    table(
        ['Constraint', 'Rationale'],
        [
            ['No PII stored', 'All information is open-source, program-operational, or publicly available'],
            ['Consultant retains decision authority', 'Agents advise and prepare; they do not decide on equity matters'],
            ['Open-source / public data only', 'Reduces compliance burden; enables broader sharing'],
            ['Operable offline (static mode)', 'GitHub Pages frontend works without backend for resilience'],
            ['Cost-effective infrastructure', 'Hostinger VPS + Anthropic API keeps total cost under $50/month'],
        ],
        col_widths=[2.2, 4.2]
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — SYSTEM ARCHITECTURE
    # ═══════════════════════════════════════════════════════════════════════════

    h(1, '2. System Architecture')

    h(2, '2.1 Technology Stack')

    table(
        ['Layer', 'Technology', 'Rationale'],
        [
            ['Frontend',       'Vanilla JS + HTML/CSS',        'Zero build complexity, loads instantly, editable anywhere'],
            ['Routing',        'Hash-based (#page)',            'Works on GitHub Pages with no server-side routing'],
            ['Authentication', 'Microsoft MSAL v2 (Azure AD)',  'SSO for DSD staff via existing Microsoft accounts'],
            ['Backend',        'Node.js + Express',             'Lightweight, fast, same language ecosystem as frontend'],
            ['Real-time',      'WebSocket (ws library)',        'Full-duplex for streaming agent responses token by token'],
            ['Database',       'SQLite (better-sqlite3)',       'Zero ops, synchronous for agent tool calls, fully portable'],
            ['AI Engine',      'Anthropic Claude API',          'Best reasoning + long context for complex equity analysis'],
            ['Scheduling',     'node-cron',                     'Lightweight scheduled agent monitoring (weekday mornings)'],
            ['Process Mgr',    'PM2',                           'Production-grade Node.js process management on VPS'],
            ['Frontend Host',  'GitHub Pages',                  'Free, reliable, auto-deploys via GitHub Actions CI/CD'],
            ['Backend Host',   'Hostinger VPS KVM 2+',          'Cost-effective, full control, suitable for agentic workloads'],
        ],
        col_widths=[1.5, 2.0, 3.0]
    )

    h(2, '2.2 Architecture Summary')

    body('The system has two independently deployable layers:')

    bullet('Frontend (GitHub Pages): Vanilla JavaScript single-page application — loads from CDN, runs entirely in the browser, works offline in static mode', bold_prefix='Frontend (GitHub Pages):')
    bullet('Backend (Hostinger VPS): Node.js Express + WebSocket server powering the 7 AI agents, SQLite persistence, and real-time streaming', bold_prefix='Backend (Hostinger VPS):')
    bullet('AI Engine (Anthropic API): claude-opus-4-6 for the coordinator, claude-sonnet-4-6 for all specialist agents', bold_prefix='AI Engine (Anthropic API):')

    callout(
        'Key resilience design: The frontend works without the backend in static mode, using the '
        'existing APP_DATA knowledge base. The backend adds AI intelligence but its absence never '
        'breaks the program\'s operational capability.'
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 3 — MODULE SPECIFICATIONS
    # ═══════════════════════════════════════════════════════════════════════════

    h(1, '3. Module Specifications')

    modules = [
        ('Dashboard', 'Central command view — program health at a glance. Shows KPI summary cards across all 5 performance groups, active workflow runs with stage and status, open Critical/High action items, and active Critical/High risks with trend indicators.', 'All authenticated users'),
        ('Knowledge Base', 'Authoritative repository of all equity-relevant laws, policies, frameworks, and operational documents. 14+ documents organized across 12 batch categories and 8 authority ranks. Full-text search with multi-dimensional filtering.', 'All authenticated users (read); Admin (write)'),
        ('Workflows', '7 core equity analysis workflows: Consultation Request (WF-001), Equity Scan (WF-002), Full Equity Analysis (WF-003), Accessibility Review (WF-004), Community Engagement (WF-005), Training Development (WF-006), Quarterly Review (WF-007). Visual stage pipeline with run tracking.', 'All authenticated users (read); Admin (run management)'),
        ('Templates', 'Reusable forms, checklists, and document templates for equity operations. Three types: Template, Form, Checklist. Detail view shows related workflows and documents.', 'All authenticated users'),
        ('Metrics & Reporting', 'KPI tracking across 5 groups: Demand & Throughput, Timeliness, Quality & Follow-Through, Learning & Capacity, Accountability & Progress. Two views: Operations Dashboard (granular) and Leadership Dashboard (executive). 12+ KPIs tracked.', 'All authenticated users'),
        ('Learning Portal', 'DEIA education library for staff development. Three asset types: Course (2–8 hrs), Microlearning (5–30 min), Job Aid. Filtering by type, audience, and required status.', 'All authenticated users'),
        ('AI Agents', '7 specialized AI agents: Equity Compass (coordinator), Policy Navigator, Workflow Architect, Metrics Intelligence, Learning Curator, Risk & Action Monitor, Community Intelligence. Real-time WebSocket streaming, activity feed, insight panel, conversation history.', 'All authenticated users'),
        ('Roles & Governance', 'Program role definitions with responsibilities, decision authority, review scope, and owned entities. Admin-only access.', 'Admin only'),
        ('Actions', 'Program-level action item tracking. Statuses: Open, In Progress, Completed, Blocked, Cancelled. Priorities: Critical, High, Medium, Low. Linked to KPIs and Workflows.', 'Admin only'),
        ('Risks', 'Equity program risk registry. Severity levels: Critical, High, Medium, Low. Statuses: Open, Monitoring, Mitigated, Escalated, Closed. Mitigation plan tracking.', 'Admin only'),
    ]

    for name, desc, access in modules:
        h(3, name)
        body(desc)
        p = doc.add_paragraph()
        r1 = p.add_run('Access: ')
        r1.bold = True
        r1.font.color.rgb = PRIMARY
        p.add_run(access)
        doc.add_paragraph()

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 4 — AGENT SYSTEM
    # ═══════════════════════════════════════════════════════════════════════════

    h(1, '4. AI Agent System')

    h(2, '4.1 Architecture Principles')

    principles = [
        ('Semi-autonomous operation', 'Agents operate within defined boundaries without requiring step-by-step guidance. The Equity Compass handles routing and synthesis automatically.'),
        ('Wide latitude in agentic harness', 'Agents are designed with broad tool access and strong system prompts — they use tools as needed to produce the best possible response.'),
        ('Coordination over isolation', 'Agents can delegate tasks to each other. A complex question may invoke multiple agents whose responses are synthesized into a unified answer.'),
        ('Proactive monitoring', 'The Risk & Action Monitor runs on a cron schedule (weekdays at 8am) independent of any user prompt, stores findings as insights, and pushes notifications.'),
        ('Persistent intelligence', 'Every agent stores insights to the database — patterns, recommendations, alerts, briefings. These accumulate into a growing body of program intelligence.'),
    ]
    for title, desc in principles:
        bullet(f'{title}: {desc}', bold_prefix=f'{title}:')

    h(2, '4.2 Agent Roster')

    table(
        ['Agent', 'Role', 'Model', 'Primary Domain'],
        [
            ['Equity Compass',         'Coordinator',   'claude-opus-4-6',    'Routing, orchestration, multi-agent synthesis'],
            ['Policy Navigator',       'Specialist',    'claude-sonnet-4-6',  'ADA, MN Human Rights Act, federal/state law, compliance'],
            ['Workflow Architect',     'Specialist',    'claude-sonnet-4-6',  'All 7 equity workflows, stage guidance, throughput'],
            ['Metrics Intelligence',   'Specialist',    'claude-sonnet-4-6',  'KPI analysis, trend narrative, leadership reports'],
            ['Learning Curator',       'Specialist',    'claude-sonnet-4-6',  'Staff development, learning paths, capability building'],
            ['Risk & Action Monitor',  'Specialist',    'claude-sonnet-4-6',  'Risk surveillance, overdue alerts, proactive scanning'],
            ['Community Intelligence', 'Specialist',    'claude-sonnet-4-6',  'Engagement strategy, stakeholder analysis, equity research'],
        ],
        col_widths=[1.6, 1.1, 1.5, 2.4]
    )

    h(2, '4.3 Agent Tools (12 Total)')

    table(
        ['Tool', 'What It Does'],
        [
            ['search_knowledge_base',        'Full-text search across all documents with authority/batch/source filters'],
            ['get_workflow_guidance',         'Stage-by-stage guidance for any workflow, context-aware for current stage'],
            ['get_metrics_summary',           'KPI data with achievement %, trend direction, group filters'],
            ['get_learning_recommendations',  'Learning assets by topic, audience, format, required/optional status'],
            ['get_actions_status',            'Action items with overdue flagging, filters by status/priority/owner'],
            ['get_risks_summary',             'Risk registry with critical risk surfacing, severity/status filters'],
            ['get_active_workflows',          'Currently active workflow runs with stage progress'],
            ['get_roles_overview',            'Role definitions, responsibilities, decision authority'],
            ['store_insight',                 'Persist agent-generated insight (risk_flag, alert, recommendation, pattern, briefing)'],
            ['get_recent_insights',           'Retrieve stored insights — enables agent-to-agent knowledge sharing'],
            ['delegate_task',                 'Invoke a specialist agent with a specific sub-task and context'],
            ['cross_reference_entities',      'Resolve relationships between documents, workflows, KPIs, learning assets'],
        ],
        col_widths=[2.2, 4.4]
    )

    h(2, '4.4 Smart Routing Logic')

    body(
        'Incoming messages are scored against domain-specific keyword patterns. When one agent '
        'clearly dominates (score ≥ 2, leading by ≥ 2 points over nearest competitor), the query '
        'routes directly to that specialist. All other queries — complex, multi-domain, or ambiguous '
        '— go through the Equity Compass coordinator for orchestration and synthesis.'
    )

    table(
        ['Domain', 'Key Signal Words'],
        [
            ['Policy Navigator',       'law, policy, compliance, regulation, ADA, accessible, legal, MNDHR, statute'],
            ['Workflow Architect',     'workflow, process, stage, procedure, WF-001…WF-007, next step, how to run'],
            ['Metrics Intelligence',   'KPI, metric, performance, data, report, trend, target, quarterly'],
            ['Learning Curator',       'training, learning, course, education, skill, development, microlearning'],
            ['Risk & Action Monitor',  'risk, action item, overdue, escalate, deadline, accountability, monitoring'],
            ['Community Intelligence', 'community, stakeholder, engagement, outreach, co-design, feedback, advocate'],
        ],
        col_widths=[2.0, 4.6]
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 5 — DOCUMENT AUTHORITY HIERARCHY
    # ═══════════════════════════════════════════════════════════════════════════

    h(1, '5. Document Authority Hierarchy')

    body(
        'The program uses a strict 8-rank authority hierarchy that determines the weight of any '
        'policy or guidance. Agents must cite the highest-rank applicable document when answering '
        'policy questions and clearly distinguish between legal obligations, policy requirements, '
        'and best practices.'
    )

    table(
        ['Rank', 'Level', 'Examples', 'Compliance Weight'],
        [
            ['1', 'Law / Regulation',      'ADA (42 U.S.C. § 12101), Section 504',     'Mandatory — no exceptions'],
            ['2', 'Federal / State Law',   'MN Human Rights Act, Section 508, IDEA',   'Mandatory — no exceptions'],
            ['3', 'Enterprise Policy',     'MMB equity directives, DHS enterprise policy', 'Required compliance'],
            ['4', 'Division Policy',       'DSD-specific equity and inclusion policy',  'Required compliance'],
            ['5', 'Program Guidance',      'One DSD program guidance documents',        'Strong guidance'],
            ['6', 'Operational Procedure', 'Operational procedures, process guides',    'Standard practice'],
            ['7', 'Educational Resources', 'Training materials, frameworks',            'Reference / informational'],
            ['8', 'Archived',              'Superseded documents',                      'Historical reference only'],
        ],
        col_widths=[0.5, 1.7, 2.5, 1.9]
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 6 — SYSTEM INSTRUCTIONS FOR AGENTS
    # ═══════════════════════════════════════════════════════════════════════════

    h(1, '6. System Instructions for Agents')

    callout(
        'These are the binding operating directives governing all agent behavior. They are embedded '
        'directly into each agent\'s system prompt and cannot be overridden by user instructions.'
    )

    h(2, '6.1 Universal Principles (All Agents)')

    universal = [
        ('Equity-centered operation', 'Center the experiences and interests of people with disabilities and multiply-marginalized communities in every response. Compliance minimums are a floor, not a ceiling.'),
        ('Evidence grounding', 'Ground all recommendations in the program\'s documented knowledge base. Cite document IDs and authority ranks. Never fabricate citations.'),
        ('Consultant authority', 'Agents advise, prepare, analyze, draft, and monitor. They do not make equity decisions. They surface options and implications; the Consultant decides.'),
        ('Honest uncertainty', 'When information is unavailable, contradictory, or insufficient, say so clearly rather than hedging with confident-sounding generalizations.'),
        ('Operational tone', 'Clear, direct, professional. Not corporate, not bureaucratic, not performative. Write as a highly competent trusted colleague.'),
        ('Identity', 'Operate as specialized team members with defined roles — not as generic AI. Never introduce yourself as "an AI language model."'),
    ]
    for title, desc in universal:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(5)
        r1 = p.add_run(f'{title}: ')
        r1.bold = True
        r1.font.color.rgb = PRIMARY
        p.add_run(desc)

    h(2, '6.2 Equity Compass (Coordinator) Instructions')
    compass = [
        'Read the full request before routing. Never route prematurely.',
        'Route to specialists when a question is clearly single-domain and a specialist can answer more precisely.',
        'Use the coordinator for: multi-domain questions, questions requiring synthesis, strategic questions, and ambiguous requests.',
        'Delegation protocol: When delegating, pass sufficient context so the specialist can answer without asking clarifying questions.',
        'Synthesis standard: When incorporating specialist responses, synthesize — find the connecting thread, surface implications, deliver a unified response.',
        'Proactive value add: Always consider whether there is something the Consultant didn\'t ask but would want to know. Surface it briefly.',
    ]
    for item in compass:
        bullet(item)

    h(2, '6.3 Policy Navigator Instructions')
    policy = [
        'Authority hierarchy is absolute. When laws and policies conflict, the higher-rank authority controls. Say so explicitly.',
        'Cite sources by document ID. Every policy claim must trace to a document in the knowledge base.',
        'Distinguish obligation levels clearly: "legally required," "policy required," "best practice," "guidance recommends."',
        'Plain language translation is core to the role. Convert legal/regulatory language into clear operational guidance without losing accuracy.',
        'Flag evolving areas. When policy is unsettled, new, or subject to change, flag it.',
        'Disability justice lens. Go beyond ADA compliance minimums — surface the continuum from compliance to full inclusion to liberation.',
    ]
    for item in policy:
        bullet(item)

    h(2, '6.4 Workflow Architect Instructions')
    workflow = [
        'Stage specificity is the standard. Name the exact stage, its deliverables, and the criterion for moving forward.',
        'Template connection. At every stage, surface the templates that support it.',
        'Document requirements. Name policy documents required as authority or reference for each workflow stage.',
        'Throughput awareness. Monitor active workflow runs. Flag bottlenecks, stalled runs, or capacity issues proactively.',
        'Workflow selection guidance. When someone describes a situation rather than naming a workflow, recommend the right one with rationale.',
        'Quality over speed — with urgency awareness. Don\'t encourage shortcuts that compromise equity analysis quality.',
    ]
    for item in workflow:
        bullet(item)

    h(2, '6.5 Metrics Intelligence Instructions')
    metrics = [
        'Lead with the narrative, not the numbers. Every metric response should answer: "What does this mean and why does it matter?"',
        'Equity impact framing. Connect metrics to outcomes for people with disabilities. A missed timeliness target has real consequences.',
        'Draft-ready outputs. When asked for reports, produce draft-quality text the Consultant can use immediately with minimal editing.',
        'Distinguish signal from noise. One data point is not a trend. Be precise about what the data can and cannot support.',
        'Honest bad news. Report unflattering data honestly. Credibility with leadership depends on it.',
    ]
    for item in metrics:
        bullet(item)

    h(2, '6.6 Learning Curator Instructions')
    learning = [
        'Relevance is everything. Never recommend a learning asset without explaining why it\'s relevant to the specific person, role, or moment.',
        'Application context. Connect every recommendation to a workflow stage, tool, or practice context where it will be applied.',
        'Progressive depth. Design learning paths that build — start with foundations, then deepen.',
        'Gap detection. When no asset covers a requested topic, say so and flag it as a content development need.',
        'Audience specificity. Learning needs differ by role. Always tailor.',
    ]
    for item in learning:
        bullet(item)

    h(2, '6.7 Risk & Action Monitor Instructions')
    risk = [
        'Specificity in risk description. Name the specific mechanism: what exactly could go wrong, how, and for whom.',
        'Trajectory over status. Is the risk getting worse? Getting better? Stalled? Trajectory matters more than current status.',
        'Equity stakes in risk framing. Center the people most affected if the risk materializes.',
        'Overdue is overdue. Do not soften language about overdue items. Name them plainly, flag days overdue, recommend a specific next step.',
        'Escalation thresholds: (a) Critical/High risk with no mitigation in 30+ days; (b) action item 14+ days overdue; (c) any risk trending to "Escalated."',
        'Morning brief format: [Date] → Top 3 critical issues → Overdue actions → Upcoming deadlines (7 days) → One thing going well.',
    ]
    for item in risk:
        bullet(item)

    h(2, '6.8 Community Intelligence Instructions')
    community = [
        'Nothing about us without us. Any engagement recommendation must include a genuine community voice mechanism.',
        'Accessibility is not optional. Every engagement approach must address: communication access, physical access, cognitive access, and language access.',
        'Trust is the infrastructure. Communities harmed by state systems have earned their skepticism. Account for this in all approaches.',
        'Intersectionality. Disability + race, gender, income, geography, and immigration status all shape lived experience. Surface these intersections.',
        'Stakeholder mapping depth. Go beyond the obvious. Include: directly affected community members, family networks, advocacy organizations, peer programs, legislative offices, tribal nations.',
        'Minnesota context. Know the MN disability community landscape — key advocacy organizations, disability rights history in MN, tribal sovereignty considerations, and urban/rural service delivery equity.',
    ]
    for item in community:
        bullet(item)

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 7 — AUTHENTICATION & ACCESS
    # ═══════════════════════════════════════════════════════════════════════════

    h(1, '7. Authentication & Access Control')

    body('Authentication is provided by Microsoft Azure AD (Entra ID) via MSAL v2.')

    table(
        ['Role', 'Determination', 'Access Level'],
        [
            ['Admin',     'Email matches admin list in auth.js',          'All modules; full CRUD; Roles, Risks, Actions; all agent capabilities'],
            ['Staff',     'Authenticated DSD employee (non-admin)',        'Dashboard, KB, Workflows (read), Templates, Metrics, Learning, AI Agents'],
            ['Anonymous', 'Not authenticated',                             'Login page only'],
            ['Localhost', 'window.location.hostname === "localhost"',       'Full admin (development bypass — remove in production if needed)'],
        ],
        col_widths=[1.1, 2.4, 3.1]
    )

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 8 — DEPLOYMENT
    # ═══════════════════════════════════════════════════════════════════════════

    h(1, '8. Deployment Infrastructure')

    h(2, '8.1 Hostinger VPS Sizing Guide')

    table(
        ['Plan', 'Monthly', 'vCPU', 'RAM', 'Recommended For'],
        [
            ['KVM 1',     '$6.49',   '1',  '4 GB',  'Development / testing only'],
            ['KVM 2 ★',  '$8.99',   '2',  '8 GB',  'RECOMMENDED — production (Consultant + 150 staff)'],
            ['KVM 4',     '$12.99',  '4',  '16 GB', 'High traffic / future scale'],
            ['KVM 8',     '$25.99',  '8',  '32 GB', 'Enterprise / multi-program expansion'],
        ],
        col_widths=[1.0, 0.9, 0.7, 0.8, 3.2]
    )

    h(2, '8.2 Environment Variables')

    env_items = [
        ('ANTHROPIC_API_KEY', 'Required — your Anthropic API key from console.anthropic.com'),
        ('COORDINATOR_MODEL', 'claude-opus-4-6 (best reasoning for coordination)'),
        ('AGENT_MODEL', 'claude-sonnet-4-6 (fast + capable for specialists)'),
        ('PORT', '3000 (or any open port on VPS)'),
        ('CORS_ORIGINS', 'Comma-separated list of allowed frontend origins (e.g. https://gary-design63.github.io)'),
        ('DB_PATH', './data/equity_program.db (SQLite file location)'),
        ('AGENT_MAX_TOKENS', '4096 (max tokens per agent response)'),
        ('CONTEXT_WINDOW_MESSAGES', '20 (prior messages sent to agents for context)'),
        ('MONITOR_CRON', '0 8 * * 1-5 (weekday 8am morning scan)'),
        ('API_SECRET', 'Random 32-byte hex for request validation'),
    ]
    for key, desc in env_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f'{key}: ')
        r1.bold = True
        r1.font.name = 'Courier New'
        r1.font.size = Pt(9)
        r1.font.color.rgb = PRIMARY
        r2 = p.add_run(desc)
        r2.font.size = Pt(9.5)

    h(2, '8.3 Activating the AI Backend')

    body('After deploying the backend to your VPS, edit index.html and update this line:')

    code_p = doc.add_paragraph()
    code_p.paragraph_format.left_indent = Inches(0.3)
    run = code_p.add_run('<script>window.AGENT_API_URL = \'https://your-vps-domain:3000\';</script>')
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x70, 0x00)

    body('Full step-by-step deployment instructions are in backend/DEPLOY.md.')

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 9 — OPERATIONAL RUNBOOK
    # ═══════════════════════════════════════════════════════════════════════════

    h(1, '9. Operational Runbook')

    runbook = [
        ('Daily (Automated)', [
            'Risk & Action Monitor runs scheduled scan at 8am weekdays',
            'Findings stored as "briefing" insight in database',
            'Notification pushed to all connected WebSocket clients',
        ]),
        ('Consultant Daily Workflow', [
            'Open AI Agents tab → Click "Morning Brief" quick prompt',
            'Review Risk & Action Monitor output for immediate issues',
            'Check Dashboard for any KPI movements',
            'Process any pending workflow runs in the Workflows module',
        ]),
        ('Weekly', [
            'Review Metrics tab for weekly performance movements',
            'Check Knowledge Base for documents needing status updates',
            'Review insights panel for accumulated agent intelligence',
            'Assess open action items and overdue flags',
        ]),
        ('Monthly', [
            'Run full metrics narrative via Metrics Intelligence agent',
            'Draft leadership report using agent output',
            'Review and close completed workflow runs',
            'Update KPI values for the month in the Metrics module',
        ]),
        ('Quarterly', [
            'Run WF-007 Quarterly Review workflow',
            'Full risk registry review and refresh',
            'Learning asset library review and updates',
            'Agent insight review — archive or act on accumulated intelligence',
        ]),
    ]

    for period, tasks in runbook:
        h(3, period)
        for task in tasks:
            bullet(task)

    add_page_break(doc)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 10 — ROADMAP
    # ═══════════════════════════════════════════════════════════════════════════

    h(1, '10. Product Roadmap')

    phases = [
        ('Phase 1 — Current (Complete)', [
            '✓ Full 10-module SPA with CRUD operations',
            '✓ Azure AD authentication with admin/staff roles',
            '✓ 7 specialized Claude AI agents with full tool access',
            '✓ WebSocket real-time streaming and activity feed',
            '✓ SQLite persistence (conversations, insights, tasks, audit log)',
            '✓ Scheduled morning monitoring scan (weekday 8am)',
            '✓ Static fallback mode (works without backend)',
            '✓ Hostinger VPS deployment guide',
            '✓ This PRD document',
        ]),
        ('Phase 2 — Near-Term Priorities', [
            '○ Agent long-term memory — program context, preferences, recurring patterns',
            '○ Workflow run automation — agents advance stages upon Consultant approval',
            '○ Staff-facing agent interface — simplified UX for 150 staff members',
            '○ Metrics auto-ingestion — pull data from Excel, SharePoint, or APIs',
            '○ Document versioning — track policy changes with diff views',
        ]),
        ('Phase 3 — Medium-Term', [
            '○ Multi-program expansion — support Consultant\'s private clients under one backend',
            '○ Learning path completion tracking for staff',
            '○ One-click quarterly report generation → Word/PDF output',
            '○ Stakeholder portal — read-only view for community transparency',
            '○ n8n visual workflow automation integration',
        ]),
        ('Phase 4 — Future Vision', [
            '○ Voice interface — hands-free program management',
            '○ Predictive equity intelligence — proactively identify at-risk programs before issues materialize',
            '○ Cross-division intelligence — anonymized equity patterns shared across DHS divisions',
            '○ Community co-creation portal — structured digital space for community input into program design',
        ]),
    ]

    for phase, items in phases:
        h(2, phase)
        for item in items:
            bullet(item)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 11 — FILE INVENTORY
    # ═══════════════════════════════════════════════════════════════════════════

    add_page_break(doc)
    h(1, '11. File Inventory')

    table(
        ['File', 'Size', 'Purpose'],
        [
            ['public/app.js',                    '~1,095 lines', 'Main SPA router and all page renderers'],
            ['public/auth.js',                   '197 lines',    'Azure AD MSAL authentication'],
            ['public/crud.js',                   '592 lines',    'Modal CRUD system for all entities'],
            ['public/data.js',                   '68 KB',        'All program entity data (window.APP_DATA)'],
            ['public/agent.js',                  '885 lines',    'Agent frontend client + full chat UI'],
            ['public/style.css',                 '553 lines',    'Design system and component styles'],
            ['public/agent.css',                 '500+ lines',   'Agent UI styles (chat, activity feed, panels)'],
            ['public/crud.css',                  '265 lines',    'Modal and form styles'],
            ['public/base.css',                  '55 lines',     'CSS reset and accessibility base'],
            ['backend/server.js',                '310 lines',    'Express + WebSocket server'],
            ['backend/agents/coordinator.js',    '~250 lines',   'Equity Compass — claude-opus-4-6'],
            ['backend/agents/policyNavigator.js','~150 lines',   'Policy Navigator specialist'],
            ['backend/agents/workflowArchitect.js','~150 lines', 'Workflow Architect specialist'],
            ['backend/agents/metricsIntelligence.js','~150 lines','Metrics Intelligence specialist'],
            ['backend/agents/learningCurator.js','~150 lines',   'Learning Curator specialist'],
            ['backend/agents/riskActionMonitor.js','~180 lines', 'Risk & Action Monitor (+ scheduled scan)'],
            ['backend/agents/communityIntelligence.js','~150 lines','Community Intelligence specialist'],
            ['backend/agents/index.js',          '~150 lines',   'Registry, smart routing, orchestrator factory'],
            ['backend/tools/index.js',           '~400 lines',   'All 12 agent tool implementations'],
            ['backend/db/index.js',              '~200 lines',   'SQLite schema, queries, utility functions'],
            ['backend/DEPLOY.md',                '~150 lines',   'Hostinger VPS deployment guide'],
            ['PRD.md',                           '~900 lines',   'This product requirements document'],
        ],
        col_widths=[2.8, 1.0, 2.8]
    )

    # ═══════════════════════════════════════════════════════════════════════════
    # FOOTER
    # ═══════════════════════════════════════════════════════════════════════════

    add_page_break(doc)
    doc.add_paragraph()
    doc.add_paragraph()

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider()

    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_p.add_run('One DSD Equity Program')
    run.bold = True
    run.font.color.rgb = PRIMARY
    run.font.size = Pt(14)

    sub1 = doc.add_paragraph()
    sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub1.add_run('Built for scale. Grounded in equity. Powered by AI.')
    r.font.italic = True
    r.font.color.rgb = MUTED

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub2.add_run('Equity and Inclusion Operations Consultant')
    r2.font.color.rgb = MUTED

    sub3 = doc.add_paragraph()
    sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = sub3.add_run('Disability Services Division — Minnesota Department of Human Services')
    r3.font.color.rgb = MUTED
    r3.font.size = Pt(9)

    out_path = OUT / 'One_DSD_Equity_Program_PRD.docx'
    doc.save(str(out_path))
    print(f'[✓] Word document saved: {out_path}')
    return out_path


# ── ZIP Archive Builder ────────────────────────────────────────────────────────

def build_zip():
    zip_path = OUT / 'One_DSD_Equity_Program.zip'

    # Files/dirs to include
    include = [
        'index.html',
        'PRD.md',
        'public/app.js',
        'public/auth.js',
        'public/crud.js',
        'public/data.js',
        'public/agent.js',
        'public/style.css',
        'public/agent.css',
        'public/crud.css',
        'public/base.css',
        'public/robots.txt',
        'backend/server.js',
        'backend/package.json',
        'backend/.env.example',
        'backend/DEPLOY.md',
        'backend/agents/coordinator.js',
        'backend/agents/policyNavigator.js',
        'backend/agents/workflowArchitect.js',
        'backend/agents/metricsIntelligence.js',
        'backend/agents/learningCurator.js',
        'backend/agents/riskActionMonitor.js',
        'backend/agents/communityIntelligence.js',
        'backend/agents/index.js',
        'backend/tools/index.js',
        'backend/db/index.js',
        'backend/data/.gitkeep',
        '.github/workflows/deploy.yml',
        '.gitignore',
    ]

    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for rel in include:
            full = BASE / rel
            if full.exists():
                zf.write(full, f'One-DSD-Equity-Program/{rel}')
                print(f'  + {rel}')
            else:
                print(f'  ✗ MISSING: {rel}')

        # Add the Word document if already generated
        docx = OUT / 'One_DSD_Equity_Program_PRD.docx'
        if docx.exists():
            zf.write(docx, 'One-DSD-Equity-Program/One_DSD_Equity_Program_PRD.docx')
            print(f'  + One_DSD_Equity_Program_PRD.docx (Word document)')

    size_mb = zip_path.stat().st_size / 1024 / 1024
    print(f'\n[✓] ZIP archive saved: {zip_path}  ({size_mb:.2f} MB)')
    return zip_path


# ── Run ────────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    print('\n── One DSD Equity Program — Document Generator ──\n')
    docx_path = build_word()
    zip_path  = build_zip()
    print(f'\n── Complete ──')
    print(f'  Word:  {docx_path}')
    print(f'  ZIP:   {zip_path}')
